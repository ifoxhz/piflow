"""Convert docs/user-manual.zh.md → Word (.docx)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "user-manual.zh.md"
OUT = ROOT / "docs" / "user-manual.zh.docx"


def set_run_font(run, east_asia: str = "微软雅黑", latin: str = "Calibri", size: Pt | None = None):
    run.font.name = latin
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = size


def add_inline(paragraph, node: Tag | NavigableString, *, bold=False, italic=False, code=False):
    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        run = paragraph.add_run(text)
        set_run_font(run, size=Pt(10.5))
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = "Consolas"
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        return

    if not isinstance(node, Tag):
        return

    name = node.name.lower()
    if name in {"strong", "b"}:
        for child in node.children:
            add_inline(paragraph, child, bold=True, italic=italic, code=code)
        return
    if name in {"em", "i"}:
        for child in node.children:
            add_inline(paragraph, child, bold=bold, italic=True, code=code)
        return
    if name in {"code"}:
        for child in node.children:
            add_inline(paragraph, child, bold=bold, italic=italic, code=True)
        return
    if name == "a":
        text = node.get_text()
        run = paragraph.add_run(text)
        set_run_font(run, size=Pt(10.5))
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        run.underline = True
        return
    if name == "br":
        paragraph.add_run().add_break()
        return

    for child in node.children:
        add_inline(paragraph, child, bold=bold, italic=italic, code=code)


def add_paragraph_from_tag(doc: Document, tag: Tag, *, style: str | None = None):
    p = doc.add_paragraph(style=style)
    for child in tag.children:
        add_inline(p, child)
    return p


def add_table(doc: Document, table_tag: Tag):
    rows = table_tag.find_all("tr")
    if not rows:
        return
    grid = []
    for tr in rows:
        cells = tr.find_all(["th", "td"])
        grid.append([c.get_text(strip=True) for c in cells])
    cols = max(len(r) for r in grid)
    table = doc.add_table(rows=len(grid), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(grid):
        for j in range(cols):
            text = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            set_run_font(run, size=Pt(9))
            if i == 0 or rows[i].find("th"):
                run.bold = True


def convert_html(doc: Document, soup: BeautifulSoup):
    body = soup.body or soup
    for el in body.children:
        if isinstance(el, NavigableString):
            continue
        if not isinstance(el, Tag):
            continue
        name = el.name.lower()
        if name in {"h1", "h2", "h3", "h4"}:
            level = int(name[1])
            p = doc.add_heading("", level=min(level, 3))
            # clear default run and rewrite with CJK font
            if p.runs:
                p.runs[0].text = ""
            for child in el.children:
                add_inline(p, child)
            for run in p.runs:
                set_run_font(run, size=Pt(18 if level == 1 else 14 if level == 2 else 12))
                run.bold = True
        elif name == "p":
            add_paragraph_from_tag(doc, el)
        elif name in {"ul", "ol"}:
            ordered = name == "ol"
            for i, li in enumerate(el.find_all("li", recursive=False), start=1):
                style = "List Number" if ordered else "List Bullet"
                p = doc.add_paragraph(style=style)
                for child in li.children:
                    if isinstance(child, Tag) and child.name in {"ul", "ol"}:
                        continue
                    add_inline(p, child)
        elif name == "pre":
            code = el.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code.rstrip() + "\n")
            run.font.name = "Consolas"
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.2)
        elif name == "blockquote":
            p = add_paragraph_from_tag(doc, el)
            p.paragraph_format.left_indent = Inches(0.25)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif name == "table":
            add_table(doc, el)
            doc.add_paragraph()
        elif name == "hr":
            doc.add_paragraph("─" * 32)


def strip_yaml_front_matter(text: str) -> str:
    if text.startswith("---"):
        m = re.match(r"^---\r?\n.*?\r?\n---\r?\n", text, flags=re.S)
        if m:
            return text[m.end() :]
    return text


def main() -> int:
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else SRC
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else OUT
    md = strip_yaml_front_matter(src.read_text(encoding="utf-8"))
    html = markdown.markdown(
        md,
        extensions=["tables", "fenced_code", "sane_lists", "nl2br"],
        output_format="html5",
    )
    soup = BeautifulSoup(f"<body>{html}</body>", "html.parser")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    convert_html(doc, soup)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"[docx] wrote {out} ({out.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
